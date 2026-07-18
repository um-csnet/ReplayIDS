"""
Update ER rows in Tables 5 and 6 of the manuscript with re-run numbers.
Run this AFTER all ER runs complete (extract_er_metrics.py exits 0).

Updates both:
  - IDS_Journal_ComputerNetworks_v2.1_draft.docx  (Tables 5 and 6)
  - IDS_Journal_ComputerNetworks_v2_Elsevier.docx  (Tables 4, 5, 6, 7)

Table structure (v2.1_draft.docx):
  tables[4] = Table 5 CI:   rows 6-8  → ER-Balanced CI b=1,5,10
  tables[5] = Table 6 CII:  rows 6-8  → ER-Strat CII b=1,5,10
                             rows 9-11 → ER-Balanced CII b=1,5,10

Table structure (v2_Elsevier.docx):
  tables[4] = Table 4 CI:   rows 6-8  → ER-Balanced CI b=1,5,10
  tables[5] = Table 5 CII:  rows 3-5  → ER-Strat CII b=1,5,10
                             rows 6-8  → ER-Balanced CII b=1,5,10
  tables[6] = Table 6 summary: rows 3,5,6 → ForgA/ForgF/IntrA/IntrF only (cols 3-6)
  tables[7] = Table 7 poison:  row 1   → 0% clean ER baseline (cols 1-6)
"""

import re
import os
import argparse
from docx import Document

_p = argparse.ArgumentParser()
_p.add_argument(
    "--logdir",
    default="/home/minda/synology/2026/0705_AziziAfif/experiments/lwf/results/logs",
)
_p.add_argument(
    "--manuscript-dir", default="/home/minda/synology/2026/0705_AziziAfif/manuscript"
)
_args = _p.parse_args()
LOGDIR = _args.logdir
MANUSCRIPT_DIR = _args.manuscript_dir

ORACLE_CI_ACC = 0.9997
ORACLE_CI_F1 = 0.9979
ORACLE_CII_ACC = 0.9955
ORACLE_CII_F1 = 0.9795


def parse_log(path):
    if not os.path.exists(path):
        return None
    content = open(path, errors="replace").read()
    exp_matches = re.findall(
        r"Completed experience \d+ - Overall accuracy: ([\d.]+), Macro F1: ([\d.]+)",
        content,
    )
    if not exp_matches:
        return None
    bwt_acc_m = re.search(r"Overall BWT \(Accuracy\): (-?[\d.]+)", content)
    bwt_f1_m = re.search(r"Overall BWT \(F1-Score\): (-?[\d.]+)", content)
    if not bwt_acc_m or not bwt_f1_m:
        return None
    return dict(
        overall_accs=[float(m[0]) for m in exp_matches],
        overall_f1s=[float(m[1]) for m in exp_matches],
        bwt_acc=float(bwt_acc_m.group(1)),
        bwt_f1=float(bwt_f1_m.group(1)),
        n_exp=len(exp_matches),
    )


def ci_row(d):
    acc = d["overall_accs"][-1]
    f1 = d["overall_f1s"][-1]
    return [
        acc,
        f1,
        max(0.0, -d["bwt_acc"]),
        max(0.0, -d["bwt_f1"]),
        ORACLE_CI_ACC - acc,
        ORACLE_CI_F1 - f1,
    ]


def cii_row(d):
    acc = sum(d["overall_accs"]) / d["n_exp"]
    f1 = sum(d["overall_f1s"]) / d["n_exp"]
    return [
        acc,
        f1,
        max(0.0, -d["bwt_acc"]),
        max(0.0, -d["bwt_f1"]),
        ORACLE_CII_ACC - acc,
        ORACLE_CII_F1 - f1,
    ]


def get(tag, scenario, required=True):
    d = parse_log(f"{LOGDIR}/{tag}.log")
    if d is None:
        if required:
            raise SystemExit(f"ERROR: run incomplete or log missing: {tag}")
        return None
    return ci_row(d) if scenario == "ci" else cii_row(d)


def fmt(v):
    return f"{v:.4f}"


def set_cell(cell, text):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    if cell.paragraphs:
        cell.paragraphs[0].add_run(str(text))
    else:
        cell.text = str(text)


def update_row(table, row_idx, col_start, values):
    row = table.rows[row_idx]
    for i, v in enumerate(values):
        set_cell(row.cells[col_start + i], v)


# ── Load all ER metrics ───────────────────────────────────────────────────────
bal_ci = {m: get(f"er_balanced_ci_m{m}", "ci") for m in (1, 5, 10)}
strat_ci = {
    m: get(f"er_strat_ci_m{m}", "ci", required=False) for m in (1, 5, 10)
}  # display only, no table row
strat_cii = {m: get(f"er_strat_cii_m{m}", "cii") for m in (1, 5, 10)}
bal_cii = {m: get(f"er_balanced_cii_m{m}", "cii") for m in (1, 5, 10)}

print("=== ER metrics ===")
for m in (1, 5, 10):
    print(f"  ER-Bal CI   b={m}: {[fmt(v) for v in bal_ci[m]]}")
for m in (1, 5, 10):
    row = strat_ci[m]
    print(f"  ER-Strat CI b={m}: {[fmt(v) for v in row] if row else 'MISSING'}")
for m in (1, 5, 10):
    print(f"  ER-Strat CII b={m}: {[fmt(v) for v in strat_cii[m]]}")
for m in (1, 5, 10):
    print(f"  ER-Bal CII  b={m}: {[fmt(v) for v in bal_cii[m]]}")

# ── Update v2.1_draft.docx ───────────────────────────────────────────────────
draft = os.path.join(MANUSCRIPT_DIR, "IDS_Journal_ComputerNetworks_v2.1_draft.docx")
doc = Document(draft)
t5, t6 = doc.tables[4], doc.tables[5]

print(f"\nUpdating {os.path.basename(draft)}")
for row_idx, m in [(6, 1), (7, 5), (8, 10)]:
    update_row(t5, row_idx, 2, [fmt(v) for v in bal_ci[m]])
    print(f"  T5 row{row_idx} (ER-Bal CI b={m}): updated")
for row_idx, m in [(6, 1), (7, 5), (8, 10)]:
    update_row(t6, row_idx, 2, [fmt(v) for v in strat_cii[m]])
    print(f"  T6 row{row_idx} (ER-Strat CII b={m}): updated")
for row_idx, m in [(9, 1), (10, 5), (11, 10)]:
    update_row(t6, row_idx, 2, [fmt(v) for v in bal_cii[m]])
    print(f"  T6 row{row_idx} (ER-Bal CII b={m}): updated")
doc.save(draft)

# ── Update v2_Elsevier.docx ──────────────────────────────────────────────────
elsevier = os.path.join(MANUSCRIPT_DIR, "IDS_Journal_ComputerNetworks_v2_Elsevier.docx")
doc = Document(elsevier)
t4e, t5e, t6e, t7e = doc.tables[4], doc.tables[5], doc.tables[6], doc.tables[7]

print(f"\nUpdating {os.path.basename(elsevier)}")
for row_idx, m in [(6, 1), (7, 5), (8, 10)]:
    update_row(t4e, row_idx, 2, [fmt(v) for v in bal_ci[m]])
for row_idx, m in [(3, 1), (4, 5), (5, 10)]:
    update_row(t5e, row_idx, 2, [fmt(v) for v in strat_cii[m]])
for row_idx, m in [(6, 1), (7, 5), (8, 10)]:
    update_row(t5e, row_idx, 2, [fmt(v) for v in bal_cii[m]])
# Table 6 summary: cols 3-6 (ForgA, ForgF, IntrA, IntrF) only
for row_idx, vals in [(3, bal_ci[10]), (5, strat_cii[10]), (6, bal_cii[10])]:
    update_row(t6e, row_idx, 3, [fmt(v) for v in vals[2:]])  # skip acc,f1
# Table 7 poison row1: 0% clean = ER-Bal CI b=10, cols 1-6
update_row(t7e, 1, 1, [fmt(v) for v in bal_ci[10]])
doc.save(elsevier)

print("\nDone. Both manuscript files updated.")
